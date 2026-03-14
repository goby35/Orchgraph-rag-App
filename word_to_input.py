"""
Word to Input Processor for GraphRAG System (v2.0)
Script để xử lý hàng loạt file .docx/.pdf chứa CV nhân sự NovaTech Solutions

Tích hợp:
- Docling: Convert Word/PDF sang Markdown
- EasyOCR: Xử lý file scan
- PyVi: Tách từ tiếng Việt cho PhoBERT

Fixes:
- TASK Spillover: Sử dụng re.split() dựa trên Markdown headers
- Position: Chỉ lấy chức danh chính xác
- Neo4j Label: Chỉ chứa Họ tên
- Experience Years: Regex chính xác, default 0
- Phone: Capture đủ 10-11 số
- Embedding Summary: Tối ưu cho PhoBERT
- Attitude Fix: Tự động lấp đầy từ Mục tiêu/Giới thiệu
"""

import os
import re
import json
import uuid
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime

# ============================================================================
# LIBRARY IMPORTS WITH FALLBACKS
# ============================================================================

# Docling for document conversion
try:
    from docling.document_converter import DocumentConverter
    from docling.datamodel.base_models import InputFormat
    from docling.datamodel.pipeline_options import PdfPipelineOptions
    from docling.document_converter import PdfFormatOption
    DOCLING_AVAILABLE = True
except ImportError:
    DOCLING_AVAILABLE = False
    print("Warning: docling not installed. Run: pip install docling")

# EasyOCR for scanned documents
try:
    import easyocr
    EASYOCR_AVAILABLE = True
    # Lazy initialization to avoid loading model at import
    _ocr_reader = None
except ImportError:
    EASYOCR_AVAILABLE = False
    print("Warning: easyocr not installed. Run: pip install easyocr")

# python-docx as fallback
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. Run: pip install python-docx")

# PyVi for Vietnamese tokenization
try:
    from pyvi import ViTokenizer
    PYVI_AVAILABLE = True
except ImportError:
    PYVI_AVAILABLE = False
    print("Warning: pyvi not installed. Run: pip install pyvi")

# Import preprocessing functions
try:
    from preprocess_data import clean_text, segment_text
    PREPROCESS_AVAILABLE = True
except ImportError:
    PREPROCESS_AVAILABLE = False
    import unicodedata
    
    def clean_text(text: str, lowercase: bool = False) -> str:
        if not text:
            return ""
        text = unicodedata.normalize('NFC', text)
        text = re.sub(r'\s+', ' ', text).strip()
        return text.lower() if lowercase else text
    
    def segment_text(text: str) -> str:
        if PYVI_AVAILABLE:
            return ViTokenizer.tokenize(text)
        return text


# ============================================================================
# MARKDOWN HEADER PATTERNS (Ground Truth)
# ============================================================================

# Markdown header patterns for section detection
MARKDOWN_HEADER_PATTERNS = [
    r'^#{1,3}\s*(.+?)$',           # # Header, ## Header, ### Header
    r'^\*\*(.+?)\*\*\s*$',         # **Bold Header**
    r'^__(.+?)__\s*$',             # __Underline Header__
    r'^(.+?)\n[=\-]{3,}$',         # Setext-style headers
]

# Section name mappings (Vietnamese & English)
SECTION_KEYWORDS = {
    'skill': [
        'kỹ năng', 'skill', 'tech stack', 'công nghệ', 'technical',
        'chuyên môn kỹ thuật', 'programming', 'tools', 'frameworks'
    ],
    'knowledge': [
        'học vấn', 'education', 'bằng cấp', 'degree', 'trường',
        'university', 'kiến thức', 'knowledge', 'đào tạo', 'training',
        'chứng chỉ', 'certification', 'certificate'
    ],
    'attitude': [
        'thái độ', 'attitude', 'cam kết', 'commitment', 'đam mê',
        'passion', 'mục tiêu', 'objective', 'goal', 'sở thích',
        'hobby', 'interest', 'điểm mạnh', 'strength', 'giới thiệu',
        'introduction', 'about me', 'summary', 'profile',
        'hồ sơ cá nhân', 'cá nhân', 'personal'
    ],
    'thinking': [
        'tư duy', 'thinking', 'problem solving', 'giải quyết vấn đề',
        'phương pháp', 'methodology', 'dự án', 'project', 'kinh nghiệm',
        'experience', 'thành tựu', 'achievement',
        'hồ sơ cá nhân', 'cá nhân', 'personal'
    ],
    'personal_info': [
        'thông tin cá nhân', 'personal', 'liên hệ', 'contact',
        'họ và tên', 'name', 'thông tin chung', 'general'
    ]
}

# CV start patterns for splitting multiple CVs
CV_START_PATTERNS = [
    r'^#{1,2}\s*CV\s*(nhân\s*sự|ứng\s*viên)',
    r'^#{1,2}\s*(HỒ\s*SƠ|THÔNG\s*TIN)\s*(NHÂN\s*SỰ|ỨNG\s*VIÊN)',
    r'^[-=]{5,}$',  # Separator lines
    r'^\*\*Họ\s*(và|&)?\s*tên\*\*',
    r'^#{1,2}\s*Họ\s*(và|&)?\s*tên\s*:',
    r'(?i)^Họ\s*(?:và|&)?\s*tên\s*:',  # Case-insensitive plain text boundary
]


# ============================================================================
# DOCUMENT CONVERSION (Docling + EasyOCR)
# ============================================================================

def get_ocr_reader():
    """Lazy initialization of EasyOCR reader."""
    global _ocr_reader
    if _ocr_reader is None and EASYOCR_AVAILABLE:
        _ocr_reader = easyocr.Reader(['vi', 'en'], gpu=False)
    return _ocr_reader


def convert_to_markdown_docling(file_path: str) -> str:
    """
    Convert document to Markdown using Docling.
    
    Args:
        file_path: Path to Word/PDF file
    
    Returns:
        Markdown content string
    """
    if not DOCLING_AVAILABLE:
        raise ImportError("Docling is required. Install with: pip install docling")
    
    try:
        # Configure pipeline for better OCR handling
        pipeline_options = PdfPipelineOptions()
        pipeline_options.do_ocr = True  # Enable OCR for scanned content
        
        converter = DocumentConverter(
            format_options={
                InputFormat.PDF: PdfFormatOption(pipeline_options=pipeline_options)
            }
        )
        
        result = converter.convert(file_path)
        markdown_content = result.document.export_to_markdown()
        
        return markdown_content
        
    except Exception as e:
        print(f"Docling conversion failed: {e}")
        return None


def convert_to_markdown_easyocr(file_path: str) -> str:
    """
    Convert scanned document using EasyOCR.
    
    Args:
        file_path: Path to image/scanned PDF
    
    Returns:
        Extracted text as pseudo-Markdown
    """
    if not EASYOCR_AVAILABLE:
        raise ImportError("EasyOCR is required. Install with: pip install easyocr")
    
    reader = get_ocr_reader()
    if reader is None:
        return None
    
    try:
        results = reader.readtext(file_path, detail=0)
        
        # Convert OCR results to pseudo-Markdown
        lines = []
        for text in results:
            text = text.strip()
            if not text:
                continue
            
            # Detect potential headers (ALL CAPS or specific keywords)
            if text.isupper() and len(text) < 50:
                lines.append(f"## {text}")
            elif any(kw in text.lower() for kws in SECTION_KEYWORDS.values() for kw in kws):
                lines.append(f"## {text}")
            else:
                lines.append(text)
        
        return '\n'.join(lines)
        
    except Exception as e:
        print(f"EasyOCR extraction failed: {e}")
        return None


def convert_to_markdown_docx(file_path: str) -> str:
    """
    Fallback: Convert .docx to pseudo-Markdown using python-docx.
    
    Args:
        file_path: Path to .docx file
    
    Returns:
        Pseudo-Markdown content
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx is required. Install with: pip install python-docx")
    
    doc = Document(file_path)
    lines = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # Detect heading styles
        if para.style.name.startswith('Heading'):
            level = para.style.name[-1] if para.style.name[-1].isdigit() else '2'
            lines.append(f"{'#' * int(level)} {text}")
        # Detect bold text as potential headers
        elif para.runs and all(run.bold for run in para.runs if run.text.strip()):
            lines.append(f"**{text}**")
        else:
            lines.append(text)
    
    # Process tables
    for table in doc.tables:
        lines.append("")
        for row in table.rows:
            row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                lines.append(row_text)
        lines.append("")
    
    return '\n'.join(lines)


def convert_document_to_markdown(file_path: str) -> str:
    """
    Main conversion function - tries Docling first, then fallbacks.
    
    Args:
        file_path: Path to document
    
    Returns:
        Markdown content
    """
    file_ext = Path(file_path).suffix.lower()
    
    # Try Docling first (supports PDF, DOCX, images)
    if DOCLING_AVAILABLE:
        markdown = convert_to_markdown_docling(file_path)
        if markdown:
            return markdown
    
    # Fallback for .docx
    if file_ext == '.docx' and DOCX_AVAILABLE:
        return convert_to_markdown_docx(file_path)
    
    # Fallback for images/scanned with EasyOCR
    if file_ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp'] and EASYOCR_AVAILABLE:
        return convert_to_markdown_easyocr(file_path)
    
    raise ValueError(f"Cannot convert file: {file_path}. No suitable converter available.")


# ============================================================================
# MARKDOWN SECTION PARSING (Fix TASK Spillover)
# ============================================================================

def parse_markdown_sections(markdown_text: str) -> Dict[str, str]:
    """
    Parse Markdown into sections using re.split() on headers.
    This fixes the TASK Spillover issue by strictly delimiting sections.
    
    Args:
        markdown_text: Markdown content
    
    Returns:
        Dictionary mapping section names to their content
    """
    # Combined pattern to match all header styles
    header_pattern = r'(?m)^(#{1,3}\s*.+?$|\*\*.+?\*\*\s*$|__.+?__\s*$)'
    
    # Split by headers while keeping the headers
    parts = re.split(header_pattern, markdown_text, flags=re.MULTILINE)
    
    sections = {}
    current_header = 'preamble'
    current_content = []
    
    i = 0
    while i < len(parts):
        part = parts[i].strip()
        
        if not part:
            i += 1
            continue
        
        # Check if this part is a header
        is_header = False
        header_text = None
        
        # Check for # headers
        match = re.match(r'^#{1,3}\s*(.+)$', part)
        if match:
            is_header = True
            header_text = match.group(1).strip()
        
        # Check for **bold** headers
        if not is_header:
            match = re.match(r'^\*\*(.+?)\*\*\s*$', part)
            if match:
                is_header = True
                header_text = match.group(1).strip()
        
        # Check for __underline__ headers
        if not is_header:
            match = re.match(r'^__(.+?)__\s*$', part)
            if match:
                is_header = True
                header_text = match.group(1).strip()
        
        if is_header and header_text:
            # Save previous section
            if current_content:
                sections[current_header] = '\n'.join(current_content).strip()
            
            # Start new section
            current_header = clean_text(header_text)
            current_content = []
        else:
            current_content.append(part)
        
        i += 1
    
    # Save last section
    if current_content:
        sections[current_header] = '\n'.join(current_content).strip()
    
    return sections


def classify_section_to_task(section_name: str) -> Optional[str]:
    """
    Classify a section name to TASK category.
    
    Args:
        section_name: Section header text
    
    Returns:
        TASK key ('thinking', 'attitude', 'skill', 'knowledge') or None
    """
    section_lower = section_name.lower()
    
    for task_key, keywords in SECTION_KEYWORDS.items():
        if task_key == 'personal_info':
            continue
        for keyword in keywords:
            if keyword in section_lower:
                return task_key
    
    return None


def extract_task_from_sections(sections: Dict[str, str]) -> Dict[str, Dict[str, Any]]:
    """
    Extract TASK profile from parsed sections with strict boundaries.
    
    Args:
        sections: Parsed Markdown sections
    
    Returns:
        TASK profile dictionary
    """
    task_profile = {
        'thinking': {
            'description': 'Tư duy giải quyết vấn đề',
            'content': '',
            'keywords': [],
            'source_sections': []
        },
        'attitude': {
            'description': 'Thái độ và sự cam kết',
            'content': '',
            'keywords': [],
            'source_sections': []
        },
        'skill': {
            'description': 'Kỹ năng chuyên môn',
            'content': '',
            'keywords': [],
            'source_sections': []
        },
        'knowledge': {
            'description': 'Kiến thức nền tảng',
            'content': '',
            'keywords': [],
            'source_sections': []
        }
    }
    
    for section_name, content in sections.items():
        task_key = classify_section_to_task(section_name)
        
        if task_key and content:
            # Append content with clear separation
            if task_profile[task_key]['content']:
                task_profile[task_key]['content'] += ' | '
            task_profile[task_key]['content'] += clean_text(content, lowercase=False)
            task_profile[task_key]['source_sections'].append(section_name)
    
    # Extract keywords BEFORE segmentation (keywords match on raw text)
    for task_key in task_profile:
        content = task_profile[task_key]['content']
        task_profile[task_key]['keywords'] = extract_tech_keywords(content)
    
    # Segment TASK content for PhoBERT (sau bước này KHÔNG dùng regex xóa _)
    for task_key in task_profile:
        content = task_profile[task_key]['content']
        if content:
            task_profile[task_key]['content'] = segment_text(content)
    
    # FIX: Attitude fallback - fill from Mục tiêu/Giới thiệu if empty
    if not task_profile['attitude']['content']:
        fallback_sources = ['mục tiêu', 'giới thiệu', 'objective', 'about', 'summary', 'profile']
        for section_name, content in sections.items():
            if any(src in section_name.lower() for src in fallback_sources):
                cleaned = clean_text(content, lowercase=False)
                task_profile['attitude']['keywords'] = extract_tech_keywords(cleaned)
                task_profile['attitude']['content'] = segment_text(cleaned)
                task_profile['attitude']['source_sections'].append(f"{section_name} (fallback)")
                break
    
    # FIX: Thinking fallback - also check Giới thiệu for problem-solving approach
    if not task_profile['thinking']['content']:
        for section_name, content in sections.items():
            if any(src in section_name.lower() for src in ['giới thiệu', 'about', 'summary']):
                cleaned = clean_text(content, lowercase=False)
                task_profile['thinking']['content'] = segment_text(cleaned)
                task_profile['thinking']['source_sections'].append(f"{section_name} (fallback)")
                break
    
    return task_profile


def extract_tech_keywords(text: str) -> List[str]:
    """Extract technology keywords from text."""
    if not text:
        return []
    
    tech_keywords = [
        # Programming & Frameworks
        'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'go', 'rust',
        'react', 'angular', 'vue', 'svelte', 'nextjs', 'nuxt',
        'nodejs', 'django', 'flask', 'fastapi', 'spring', 'express',
        # DevOps & Cloud
        'docker', 'kubernetes', 'k8s', 'terraform', 'ansible',
        'aws', 'azure', 'gcp', 'cloud',
        # Databases
        'sql', 'nosql', 'mongodb', 'postgresql', 'mysql', 'redis', 'elasticsearch',
        'kafka', 'rabbitmq', 'celery',
        # Tools & Practices
        'git', 'github', 'gitlab', 'ci/cd', 'jenkins',
        'devops', 'agile', 'scrum', 'kanban',
        # AI/ML/Data
        'machine learning', 'deep learning', 'ai', 'ml', 'nlp', 'computer vision',
        'tensorflow', 'pytorch', 'scikit-learn', 'pandas', 'numpy',
        'data science', 'data engineering', 'etl', 'hadoop', 'spark', 'airflow',
        'graphql', 'rest', 'api', 'microservices',
        # Management & Business (C-level, Senior profiles)
        'leadership', 'strategy', 'strategic', 'fundraising', 'management',
        'team management', 'project management', 'product management',
        'business development', 'business strategy', 'partnership',
        'marketing', 'digital marketing', 'branding', 'sales',
        'finance', 'financial', 'budgeting', 'p&l', 'revenue',
        'operations', 'supply chain', 'logistics',
        'negotiation', 'stakeholder', 'investor relations',
        'mentoring', 'coaching', 'talent acquisition', 'hr',
        'governance', 'compliance', 'risk management',
        'innovation', 'transformation', 'digital transformation',
        'startup', 'entrepreneurship', 'venture capital',
        'kpi', 'okr', 'roi', 'growth hacking',
        'ceo', 'cto', 'cfo', 'coo', 'cmo', 'vp', 'director'
    ]
    
    text_lower = text.lower()
    found = []
    
    for keyword in tech_keywords:
        if keyword in text_lower:
            found.append(keyword)
    
    return list(set(found))


# ============================================================================
# FIELD EXTRACTION (Fixed Regex)
# ============================================================================

def extract_full_name(text: str) -> Optional[str]:
    """Extract full name from CV text."""
    patterns = [
        r'(?i)(?:Họ\s*(?:và|&)?\s*tên|Name)\s*[:\-]?\s*([A-ZÀ-Ỹ][a-zà-ỹ]+(?:\s+[A-ZÀ-Ỹ][a-zà-ỹ]+){1,5})',
        r'(?i)^\*\*([A-ZÀ-Ỹ][a-zà-ỹ]+(?:\s+[A-ZÀ-Ỹ][a-zà-ỹ]+){1,5})\*\*',
        r'(?i)^#\s*([A-ZÀ-Ỹ][a-zà-ỹ]+(?:\s+[A-ZÀ-Ỹ][a-zà-ỹ]+){1,5})\s*$',
    ]
    
    for pattern in patterns:
        match = re.search(pattern, text, re.MULTILINE)
        if match:
            name = match.group(1).strip()
            # Validate: name should be 2-6 words, no numbers
            words = name.split()
            if 2 <= len(words) <= 6 and not any(char.isdigit() for char in name):
                return name
    
    return None


def extract_position(text: str, name: str = None) -> Optional[str]:
    """
    Extract position/title - FIXED to only get exact title.
    Only from line below name or after keywords.
    
    Args:
        text: CV text
        name: Extracted name to find position near it
    
    Returns:
        Position string or None
    """
    # Pattern 1: Explicit position keywords
    position_patterns = [
        r'(?i)(?:Vị\s*trí|Chức\s*vụ|Chức\s*danh|Position|Title)\s*[:\-]?\s*([^\n]+)',
        r'(?i)(?:Ứng\s*tuyển)\s*[:\-]?\s*([^\n]+)',
    ]
    
    for pattern in position_patterns:
        match = re.search(pattern, text)
        if match:
            position = clean_text(match.group(1))
            # Validate: position should be relatively short
            if position and len(position) < 100 and not any(kw in position.lower() for kw in ['email', 'phone', 'điện thoại']):
                return position
    
    # Pattern 2: Line immediately after name - catch *italic* or _italic_ position
    if name:
        name_pattern = re.escape(name)
        match = re.search(rf'{name_pattern}\s*\n+([^\n]+)', text, re.IGNORECASE)
        if match:
            potential_position = match.group(1).strip()
            # Strip Markdown italic markers: *text* or _text_
            potential_position = re.sub(r'^[\*_]+|[\*_]+$', '', potential_position).strip()
            potential_position = clean_text(potential_position)
            # Validate: should look like a job title
            if potential_position and len(potential_position) < 80:
                # Skip if it contains contact info
                if not any(kw in potential_position.lower() for kw in ['@', 'email', 'phone', 'điện thoại', 'ngày sinh', 'địa chỉ']):
                    return potential_position
    
    # Pattern 3: Italic text on line below 'Họ và tên'
    italic_after_name_patterns = [
        r'(?i)Họ\s*(?:và|&)?\s*tên\s*[:\-]?\s*[^\n]*\n+\s*\*([^\*]+)\*',
        r'(?i)Họ\s*(?:và|&)?\s*tên\s*[:\-]?\s*[^\n]*\n+\s*_([^_]+)_',
    ]
    for pattern in italic_after_name_patterns:
        match = re.search(pattern, text)
        if match:
            potential_position = clean_text(match.group(1))
            if potential_position and len(potential_position) < 80:
                if not any(kw in potential_position.lower() for kw in ['@', 'email', 'phone', 'điện thoại', 'ngày sinh', 'địa chỉ']):
                    return potential_position
    
    return None


def extract_experience_years(text: str) -> int:
    """
    Extract years of experience - FIXED to return integer.
    
    Args:
        text: CV text
    
    Returns:
        Integer years (default 0)
    """
    patterns = [
        r'(\d+)\s*(?:\+)?\s*(?:năm|years?)\s*(?:kinh\s*nghiệm|experience)?',
        r'(?:kinh\s*nghiệm|experience)\s*[:\-]?\s*(\d+)\s*(?:\+)?\s*(?:năm|years?)?',
    ]
    
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            try:
                years = int(match.group(1))
                if 0 < years <= 50:  # Sanity check
                    return years
            except ValueError:
                continue
    
    return 0  # Default


def extract_phone(text: str) -> Optional[str]:
    """
    Extract phone number - FIXED to capture 10-11 digits for Vietnam.
    
    Args:
        text: CV text
    
    Returns:
        Phone number string or None
    """
    patterns = [
        # With parentheses: (+84) format
        r'(?i)(?:Điện\s*thoại|SĐT|Phone|Tel|Mobile)\s*[:\-]?\s*\(\+84\)\s*(\d[\d\s]{8,11})',
        r'\(\+84\)\s*(\d[\d\s]{8,11})',
        # With keyword
        r'(?i)(?:Điện\s*thoại|SĐT|Phone|Tel|Mobile)\s*[:\-]?\s*(\+?84|0)(\d{9,10})',
        r'(?i)(?:Điện\s*thoại|SĐT|Phone|Tel|Mobile)\s*[:\-]?\s*(\+?84|0)[\s\.\-]?(\d{2,4})[\s\.\-]?(\d{3,4})[\s\.\-]?(\d{3,4})',
        # Standalone phone pattern
        r'(?<![0-9])(\+?84|0)(3|5|7|8|9)(\d{8})(?![0-9])',
        r'(?<![0-9])(0)(2\d)(\d{8})(?![0-9])',  # Landline
    ]
    
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            groups = match.groups()
            # Combine groups into full number
            phone = ''.join(g for g in groups if g)
            # Normalize: remove spaces, dots, dashes, parentheses
            phone = re.sub(r'[\s\.\-\(\)]', '', phone)
            
            # Handle (+84) prefix → 0...
            if phone.startswith('+84'):
                phone = '0' + phone[3:]
            elif phone.startswith('84') and len(phone) > 10:
                phone = '0' + phone[2:]
            
            # Add leading 0 if missing (from (+84) capture group)
            if not phone.startswith('0') and len(phone) in [9, 10]:
                phone = '0' + phone
            
            if len(phone) in [10, 11] and phone.startswith('0'):
                return phone
    
    return None


def extract_email(text: str) -> Optional[str]:
    """Extract email address."""
    pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    match = re.search(pattern, text)
    return match.group(0) if match else None


def extract_birth_date(text: str) -> Optional[str]:
    """Extract birth date."""
    patterns = [
        r'(?i)(?:Ngày\s*sinh|Birth\s*date|DOB)\s*[:\-]?\s*(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4})',
        r'(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{4})',
    ]
    
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            return match.group(1) if match.lastindex else match.group(0)
    
    return None


def extract_department(text: str) -> Optional[str]:
    """Extract department."""
    patterns = [
        r'(?i)(?:Phòng\s*ban|Bộ\s*phận|Department|Team)\s*[:\-]?\s*([^\n]+)',
    ]
    
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            dept = clean_text(match.group(1))
            if dept and len(dept) < 100:
                return dept
    
    return None


# ============================================================================
# EMBEDDING SUMMARY (Optimized for PhoBERT)
# ============================================================================

def build_embedding_summary(name: str, position: str, task_profile: Dict[str, Dict]) -> str:
    """
    Build optimized summary for embedding with PhoBERT.
    
    Format: Tên + Vị trí + 200 chars from each TASK section
    Runs through ViTokenizer before returning.
    
    Args:
        name: Employee name
        position: Job position
        task_profile: TASK profile dictionary
    
    Returns:
        Tokenized embedding text
    """
    parts = []
    
    # Add name and position
    if name:
        parts.append(f"Họ tên: {name}")
    if position:
        parts.append(f"Vị trí: {position}")
    
    # Add 200 chars from each TASK section
    for task_key in ['skill', 'knowledge', 'thinking', 'attitude']:
        content = task_profile.get(task_key, {}).get('content', '')
        if content:
            # Truncate to 200 chars at word boundary
            if len(content) > 200:
                truncated = content[:200]
                # Find last space to avoid cutting words
                last_space = truncated.rfind(' ')
                if last_space > 150:
                    truncated = truncated[:last_space]
                content = truncated + '...'
            parts.append(f"{task_key.upper()}: {content}")
    
    summary = ' '.join(parts)
    
    # Tokenize for PhoBERT via segment_text (handles fallback internally)
    # Sau bước này KHÔNG dùng regex xóa _ vì PhoBERT cần từ ghép: kiến_trúc_sư
    summary = segment_text(summary)
    
    return summary


# ============================================================================
# CV PROCESSING PIPELINE
# ============================================================================

def split_cvs_from_markdown(markdown_text: str) -> List[str]:
    """
    Split multiple CVs from a single Markdown document.
    
    Args:
        markdown_text: Full Markdown content
    
    Returns:
        List of individual CV texts
    """
    # Combined pattern for CV boundaries - includes 'Họ và tên:' (case-insensitive)
    cv_boundary_pattern = r'(?m)^(?:#{1,2}\s*(?:CV|HỒ\s*SƠ|THÔNG\s*TIN).+|[-=]{5,}|(?:Họ|HỌ)\s*(?:và|VÀ|&)?\s*(?:tên|TÊN)\s*:)'
    
    # Find all potential CV starts
    boundaries = list(re.finditer(cv_boundary_pattern, markdown_text, re.IGNORECASE))
    
    if not boundaries:
        # No clear boundaries - treat as single CV
        return [markdown_text.strip()] if markdown_text.strip() else []
    
    cvs = []
    for i, match in enumerate(boundaries):
        start = match.start()
        end = boundaries[i + 1].start() if i + 1 < len(boundaries) else len(markdown_text)
        
        cv_text = markdown_text[start:end].strip()
        if cv_text and len(cv_text) > 100:  # Minimum length
            cvs.append(cv_text)
    
    # If no CVs extracted, return whole text
    return cvs if cvs else [markdown_text.strip()]


def process_single_cv(cv_text: str, 
                      source_file: str,
                      cv_index: int) -> Dict[str, Any]:
    """
    Process a single CV with all fixes applied.
    
    Args:
        cv_text: CV text (Markdown format)
        source_file: Source filename
        cv_index: Index in source file
    
    Returns:
        Structured CV data
    """
    # Generate unique ID
    emp_id = f"EMP_{uuid.uuid4().hex[:8].upper()}"
    
    # Parse sections from Markdown
    sections = parse_markdown_sections(cv_text)
    
    # Clean the full text
    cleaned_text = clean_text(cv_text, lowercase=False)
    
    # Extract basic info with FIXED patterns
    full_name = extract_full_name(cleaned_text)
    position = extract_position(cleaned_text, full_name)
    experience_years = extract_experience_years(cleaned_text)
    phone = extract_phone(cleaned_text)
    email = extract_email(cleaned_text)
    birth_date = extract_birth_date(cleaned_text)
    department = extract_department(cleaned_text)
    
    # Clean extracted fields with preprocess_data (lowercase=False for PhoBERT v2)
    if full_name:
        full_name = clean_text(full_name, lowercase=False)
    if position:
        position = clean_text(position, lowercase=False)
    if department:
        department = clean_text(department, lowercase=False)
    
    # Extract TASK profile with strict boundaries (FIX: Spillover)
    # Nội dung TASK đã được clean_text + segment_text bên trong extract_task_from_sections
    task_profile = extract_task_from_sections(sections)
    
    # Build embedding summary (segment_text applied inside)
    embedding_summary = build_embedding_summary(full_name, position, task_profile)
    
    # Segment full text for storage (segment_text handles PYVI fallback internally)
    segmented_text = segment_text(cleaned_text)
    
    # Build result structure
    result = {
        'employee_id': emp_id,
        'source_file': source_file,
        'cv_index': cv_index,
        'extracted_at': datetime.now().isoformat(),
        
        # Basic Information (FIXED + no null core fields)
        'basic_info': {
            'full_name': full_name or 'Unknown',
            'position': position or 'Chưa xác định',
            'department': department or 'Chưa xác định',
            'email': email or '',
            'phone': phone or '',
            'birth_date': birth_date or '',
            'experience_years': experience_years,  # FIXED: Integer, default 0
        },
        
        # TASK Profile (FIXED: No spillover + no null content)
        'task_profile': {
            'thinking': {
                'description': task_profile['thinking']['description'],
                'content': task_profile['thinking']['content'] or '',
                'keywords': task_profile['thinking']['keywords'] or [],
            },
            'attitude': {
                'description': task_profile['attitude']['description'],
                'content': task_profile['attitude']['content'] or '',
                'keywords': task_profile['attitude']['keywords'] or [],
            },
            'skill': {
                'description': task_profile['skill']['description'],
                'content': task_profile['skill']['content'] or '',
                'keywords': task_profile['skill']['keywords'] or [],
            },
            'knowledge': {
                'description': task_profile['knowledge']['description'],
                'content': task_profile['knowledge']['content'] or '',
                'keywords': task_profile['knowledge']['keywords'] or [],
            },
        },
        
        # Text versions
        'raw_text': cv_text,
        'cleaned_text': cleaned_text,
        'segmented_text': segmented_text,
        
        # Embedding text (FIXED: Optimized summary with tokenization)
        'embedding_text': embedding_summary,
        
        # Graph metadata (FIXED: label = only name, no null)
        'graph_metadata': {
            'node_type': 'Employee',
            'label': full_name or 'Unknown',
            'properties': {
                'employee_id': emp_id,
                'position': position or 'Chưa xác định',
                'department': department or 'Chưa xác định',
                'experience_years': experience_years,
                'skills': task_profile['skill']['keywords'] or [],
                'email': email or '',
                'phone': phone or '',
            }
        },
        
        # Parsed sections for debugging
        '_debug_sections': list(sections.keys()),
    }
    
    return result


# ============================================================================
# BATCH PROCESSING
# ============================================================================

def process_folder(folder_path: str,
                   output_path: str = "clean_data.json") -> Dict[str, Any]:
    """
    Process all documents in a folder.
    
    Args:
        folder_path: Input folder path
        output_path: Output JSON file path
    
    Returns:
        Processing statistics
    """
    folder = Path(folder_path)
    
    if not folder.exists():
        print(f"⚠️  Folder not found: {folder_path}")
        print(f"   Creating folder...")
        folder.mkdir(parents=True, exist_ok=True)
        return {'error': 'Folder created, please add documents'}
    
    # Find supported files
    supported_extensions = ['.docx', '.pdf', '.doc', '.png', '.jpg', '.jpeg']
    files = []
    for ext in supported_extensions:
        files.extend(folder.glob(f"*{ext}"))
    
    if not files:
        print(f"⚠️  No supported files found in: {folder_path}")
        return {'error': 'No supported files found'}
    
    print(f"\n📁 Found {len(files)} file(s) in {folder_path}")
    print("=" * 60)
    
    all_employees = []
    stats = {
        'total_files': len(files),
        'total_cvs': 0,
        'files_processed': [],
        'errors': [],
        'conversion_method': []
    }
    
    for file_path in files:
        try:
            print(f"\n📄 Processing: {file_path.name}")
            
            # Convert to Markdown
            markdown_content = convert_document_to_markdown(str(file_path))
            
            if not markdown_content:
                raise ValueError("Failed to extract content")
            
            print(f"   ✓ Converted to Markdown ({len(markdown_content)} chars)")
            
            # Split into individual CVs
            cv_texts = split_cvs_from_markdown(markdown_content)
            print(f"   Found {len(cv_texts)} CV(s)")
            
            # Process each CV
            for i, cv_text in enumerate(cv_texts):
                try:
                    cv_data = process_single_cv(
                        cv_text=cv_text,
                        source_file=file_path.name,
                        cv_index=i + 1
                    )
                    all_employees.append(cv_data)
                    
                    name = cv_data['basic_info']['full_name']
                    position = cv_data['basic_info']['position'] or 'N/A'
                    print(f"   ✓ CV {i+1}: {name} - {position} ({cv_data['employee_id']})")
                    
                except Exception as e:
                    print(f"   ✗ Error processing CV {i+1}: {e}")
                    stats['errors'].append({
                        'file': file_path.name,
                        'cv_index': i + 1,
                        'error': str(e)
                    })
            
            stats['files_processed'].append({
                'file': file_path.name,
                'cv_count': len(cv_texts)
            })
            
            print(f"   → Đã bóc tách thành công {len(cv_texts)} CV từ file {file_path.name}")
            
        except Exception as e:
            print(f"   ✗ Error processing file: {e}")
            stats['errors'].append({
                'file': file_path.name,
                'error': str(e)
            })
    
    stats['total_cvs'] = len(all_employees)
    
    # Build output
    output_data = {
        'metadata': {
            'source_folder': str(folder_path),
            'processed_at': datetime.now().isoformat(),
            'total_employees': len(all_employees),
            'processing_stats': {
                'files_processed': len(stats['files_processed']),
                'errors': len(stats['errors']),
            },
            'converters': {
                'docling': DOCLING_AVAILABLE,
                'easyocr': EASYOCR_AVAILABLE,
                'python_docx': DOCX_AVAILABLE,
                'pyvi': PYVI_AVAILABLE,
            },
            'ready_for': ['ChromaDB', 'Neo4j', 'PhoBERT']
        },
        'employees': all_employees
    }
    
    # Save output
    print(f"\n💾 Saving to: {output_path}")
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(output_data, f, ensure_ascii=False, indent=2)
    
    return stats


def create_sample_docx(output_folder: str = "storage/cv"):
    """Create sample .docx file for testing."""
    if not DOCX_AVAILABLE:
        print("python-docx is required to create sample files")
        return None
    
    folder = Path(output_folder)
    folder.mkdir(parents=True, exist_ok=True)
    
    doc = Document()
    
    # CV 1
    doc.add_heading('CV NHÂN SỰ - NOVATECH SOLUTIONS', 0)
    doc.add_paragraph('')
    
    doc.add_heading('Nguyễn Văn An', level=1)
    doc.add_paragraph('Senior Software Engineer')
    doc.add_paragraph('')
    
    doc.add_heading('THÔNG TIN CÁ NHÂN', level=2)
    doc.add_paragraph('Email: nguyen.an@novatech.vn')
    doc.add_paragraph('Điện thoại: 0912345678')
    doc.add_paragraph('Ngày sinh: 15/03/1990')
    doc.add_paragraph('Phòng ban: Engineering')
    doc.add_paragraph('Kinh nghiệm: 5 năm')
    doc.add_paragraph('')
    
    doc.add_heading('KỸ NĂNG', level=2)
    doc.add_paragraph('Python, Java, JavaScript, React, Docker, Kubernetes, AWS')
    doc.add_paragraph('Frameworks: Django, FastAPI, Spring Boot')
    doc.add_paragraph('')
    
    doc.add_heading('HỌC VẤN', level=2)
    doc.add_paragraph('Thạc sĩ Khoa học Máy tính - Đại học Bách Khoa Hà Nội (2015)')
    doc.add_paragraph('Chứng chỉ: AWS Solutions Architect, Kubernetes Administrator')
    doc.add_paragraph('')
    
    doc.add_heading('KINH NGHIỆM LÀM VIỆC', level=2)
    doc.add_paragraph('2020-Nay: Senior Engineer tại NovaTech Solutions')
    doc.add_paragraph('Dự án tiêu biểu: Xây dựng hệ thống microservices cho e-commerce platform')
    doc.add_paragraph('Phương pháp: Agile/Scrum, System Design, Problem-solving')
    doc.add_paragraph('')
    
    doc.add_heading('MỤC TIÊU NGHỀ NGHIỆP', level=2)
    doc.add_paragraph('Trở thành Tech Lead và đóng góp vào các sản phẩm AI.')
    doc.add_paragraph('Điểm mạnh: Chủ động, sáng tạo, tinh thần teamwork cao.')
    doc.add_paragraph('')
    
    # Separator
    doc.add_paragraph('=' * 60)
    doc.add_paragraph('')
    
    # CV 2
    doc.add_heading('Trần Thị Bình', level=1)
    doc.add_paragraph('Data Scientist')
    doc.add_paragraph('')
    
    doc.add_heading('THÔNG TIN CÁ NHÂN', level=2)
    doc.add_paragraph('Email: tran.binh@novatech.vn')
    doc.add_paragraph('Điện thoại: 0987654321')
    doc.add_paragraph('Phòng ban: AI/ML Team')
    doc.add_paragraph('Kinh nghiệm: 3 năm')
    doc.add_paragraph('')
    
    doc.add_heading('KỸ NĂNG', level=2)
    doc.add_paragraph('Python, TensorFlow, PyTorch, Scikit-learn, SQL, Spark')
    doc.add_paragraph('Machine Learning, Deep Learning, NLP, Computer Vision')
    doc.add_paragraph('')
    
    doc.add_heading('HỌC VẤN', level=2)
    doc.add_paragraph('Tiến sĩ Trí tuệ Nhân tạo - KAIST, Hàn Quốc (2021)')
    doc.add_paragraph('Đào tạo: Google AI Residency')
    doc.add_paragraph('')
    
    doc.add_heading('KINH NGHIỆM', level=2)
    doc.add_paragraph('2021-Nay: Data Scientist tại NovaTech Solutions')
    doc.add_paragraph('Tư duy: Analytical thinking, Research-oriented')
    doc.add_paragraph('Giải quyết vấn đề: Áp dụng ML để tối ưu hóa quy trình kinh doanh')
    doc.add_paragraph('')
    
    doc.add_heading('GIỚI THIỆU BẢN THÂN', level=2)
    doc.add_paragraph('Đam mê nghiên cứu và ứng dụng AI vào thực tiễn.')
    doc.add_paragraph('Cam kết đóng góp vào cộng đồng AI Việt Nam.')
    
    # Save
    output_path = folder / "sample_cv_novatech.docx"
    doc.save(str(output_path))
    
    print(f"✓ Created sample file: {output_path}")
    return str(output_path)


# ============================================================================
# MAIN
# ============================================================================

def main():
    """Main entry point."""
    import argparse
    
    parser = argparse.ArgumentParser(
        description='Process CV documents for GraphRAG (v2.0 with Docling + EasyOCR)'
    )
    parser.add_argument('--input', '-i', type=str, default='storage/cv',
                        help='Input folder containing documents (default: storage/cv)')
    parser.add_argument('--output', '-o', type=str, default='clean_data.json',
                        help='Output JSON file path (default: clean_data.json)')
    parser.add_argument('--create-sample', action='store_true',
                        help='Create sample .docx file for testing')
    
    args = parser.parse_args()
    
    print("\n" + "=" * 60)
    print("  NovaTech CV Processor v2.0")
    print("  Docling + EasyOCR + PhoBERT Pipeline")
    print("=" * 60)
    
    # Print converter status
    print("\n📦 Converters Status:")
    print(f"   Docling:     {'✓ Available' if DOCLING_AVAILABLE else '✗ Not installed'}")
    print(f"   EasyOCR:     {'✓ Available' if EASYOCR_AVAILABLE else '✗ Not installed'}")
    print(f"   python-docx: {'✓ Available' if DOCX_AVAILABLE else '✗ Not installed'}")
    print(f"   PyVi:        {'✓ Available' if PYVI_AVAILABLE else '✗ Not installed'}")
    
    # Create sample if requested
    if args.create_sample:
        create_sample_docx(args.input)
    
    # Process folder
    stats = process_folder(
        folder_path=args.input,
        output_path=args.output
    )
    
    # Print summary
    print("\n" + "=" * 60)
    print("  Processing Summary")
    print("=" * 60)
    
    if 'error' in stats:
        print(f"  ⚠️  {stats['error']}")
    else:
        print(f"  Total files processed: {stats['total_files']}")
        print(f"  Total CVs extracted: {stats['total_cvs']}")
        print(f"  Errors: {len(stats['errors'])}")
        
        if stats['files_processed']:
            print("\n  Files breakdown:")
            for fp in stats['files_processed']:
                print(f"    - {fp['file']}: {fp['cv_count']} CV(s)")
        
        print(f"\n  Output: {args.output}")
        print("\n  🎯 Data ready for:")
        print("     - ChromaDB (vector embeddings - PhoBERT optimized)")
        print("     - Neo4j (employee graph - clean labels)")


if __name__ == "__main__":
    main()
